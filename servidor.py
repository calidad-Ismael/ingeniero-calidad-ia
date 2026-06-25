"""
Ingeniero de Calidad IA — Servicio de conversión y edición de documentos.

Este backend SOLO se encarga de lo que el navegador no puede hacer con fidelidad:
  1) Convertir Word (.docx) a PDF EXACTO usando LibreOffice headless
     (respeta formato, colores, layout y cantidad de hojas).
  2) Editar el contenido de un .docx preservando su formato, y subir
     automáticamente el número de revisión (texto tipo "Revisión: 01").

Todo el resto de la app (chat IA, base de conocimiento, Supabase, etc.)
corre en el navegador (app.html). Este servidor no toca Supabase ni la API
de Anthropic.
"""

import os
import re
import io
import json
import shutil
import tempfile
import subprocess
from typing import List

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document

app = FastAPI(title="Ingeniero de Calidad IA — Conversión y Edición")

# CORS abierto para que app.html (Netlify o local) pueda llamar al servicio
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Nueva-Revision", "X-Reemplazos-Aplicados", "Content-Disposition"],
)


# ---------------------------------------------------------------------------
# Utilidades
# ---------------------------------------------------------------------------
def _encontrar_soffice() -> str:
    """Devuelve la ruta del ejecutable de LibreOffice."""
    for nombre in ("soffice", "libreoffice"):
        ruta = shutil.which(nombre)
        if ruta:
            return ruta
    # rutas típicas en contenedores Linux
    for ruta in ("/usr/bin/soffice", "/usr/bin/libreoffice", "/opt/libreoffice/program/soffice"):
        if os.path.exists(ruta):
            return ruta
    raise HTTPException(status_code=500, detail="LibreOffice no está instalado en el servidor.")


def convertir_docx_a_pdf(datos_docx: bytes) -> bytes:
    """Convierte bytes de un .docx a PDF usando LibreOffice headless."""
    soffice = _encontrar_soffice()
    with tempfile.TemporaryDirectory() as tmp:
        entrada = os.path.join(tmp, "documento.docx")
        with open(entrada, "wb") as f:
            f.write(datos_docx)
        # Perfil de usuario aislado para evitar bloqueos entre conversiones concurrentes
        perfil = os.path.join(tmp, "profile")
        cmd = [
            soffice, "--headless", "--norestore", "--nolockcheck",
            "-env:UserInstallation=file://" + perfil,
            "--convert-to", "pdf", "--outdir", tmp, entrada,
        ]
        try:
            res = subprocess.run(cmd, capture_output=True, timeout=120)
        except subprocess.TimeoutExpired:
            raise HTTPException(status_code=504, detail="La conversión tardó demasiado.")
        salida = os.path.join(tmp, "documento.pdf")
        if not os.path.exists(salida):
            msg = (res.stderr or b"").decode("utf-8", "ignore")[:500]
            raise HTTPException(status_code=500, detail="No se pudo convertir a PDF. " + msg)
        with open(salida, "rb") as f:
            return f.read()


def _reemplazar_en_parrafo(parrafo, buscar: str, reemplazar: str) -> bool:
    """Reemplaza texto en un párrafo preservando el formato de los runs.

    Estrategia: primero intenta el reemplazo dentro de un único run (conserva
    el formato exacto). Si el texto buscado está partido entre varios runs,
    cae a un reemplazo a nivel de párrafo manteniendo el estilo del primer run.
    """
    cambio = False
    # 1) Reemplazo run por run (caso ideal: conserva formato)
    for run in parrafo.runs:
        if buscar in run.text:
            run.text = run.text.replace(buscar, reemplazar)
            cambio = True
    if cambio:
        return True
    # 2) Texto partido entre runs
    completo = "".join(r.text for r in parrafo.runs)
    if buscar in completo:
        nuevo = completo.replace(buscar, reemplazar)
        if parrafo.runs:
            parrafo.runs[0].text = nuevo
            for r in parrafo.runs[1:]:
                r.text = ""
        return True
    return False


def _iterar_parrafos(doc: "Document"):
    """Itera todos los párrafos del cuerpo, tablas, encabezados y pies."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for fila in t.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    yield p
    for sec in doc.sections:
        for cont in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer):
            if cont is None:
                continue
            for p in cont.paragraphs:
                yield p
            for t in cont.tables:
                for fila in t.rows:
                    for celda in fila.cells:
                        for p in celda.paragraphs:
                            yield p


_REV_PATRON = re.compile(r"(Revisi[oó]n\s*[:#]?\s*)(\d+)", re.IGNORECASE)


def _bump_revision_en_texto(texto: str):
    """Incrementa 'Revisión: 01' -> 'Revisión: 02' conservando el ancho (zero-pad)."""
    nueva = {"valor": None}

    def _sub(m):
        n = int(m.group(2))
        ancho = len(m.group(2))
        s = str(n + 1).zfill(ancho)
        nueva["valor"] = s
        return m.group(1) + s

    nuevo_texto = _REV_PATRON.sub(_sub, texto)
    return nuevo_texto, nueva["valor"]


def bump_revision(doc: "Document"):
    """Busca y sube el número de revisión en todo el documento. Devuelve la nueva."""
    nueva_rev = None
    for p in _iterar_parrafos(doc):
        completo = "".join(r.text for r in p.runs) if p.runs else p.text
        if completo and _REV_PATRON.search(completo):
            nuevo, val = _bump_revision_en_texto(completo)
            if val:
                if p.runs:
                    p.runs[0].text = nuevo
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = nuevo
                nueva_rev = val
    return nueva_rev


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------
@app.get("/")
def salud():
    try:
        soffice = _encontrar_soffice()
        ok = True
    except HTTPException:
        soffice, ok = None, False
    return {"servicio": "Ingeniero de Calidad IA - conversión/edición", "libreoffice": ok, "ruta": soffice}


@app.post("/convertir-pdf")
async def convertir_pdf(archivo: UploadFile = File(...)):
    """Recibe un .docx y devuelve el PDF EXACTO (vía LibreOffice)."""
    if not archivo.filename.lower().endswith((".docx", ".doc")):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos Word (.docx).")
    datos = await archivo.read()
    pdf = convertir_docx_a_pdf(datos)
    nombre_pdf = re.sub(r"\.docx?$", ".pdf", archivo.filename, flags=re.IGNORECASE)
    return StreamingResponse(
        io.BytesIO(pdf),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{nombre_pdf}"'},
    )


@app.post("/editar-word")
async def editar_word(
    archivo: UploadFile = File(...),
    reemplazos: str = Form("[]"),
    subir_revision: str = Form("true"),
):
    """Edita un .docx preservando formato.

    - reemplazos: JSON [{"buscar": "...", "reemplazar": "..."}]
    - subir_revision: "true"/"false" — incrementa 'Revisión: NN'
    Devuelve el .docx editado y, en la cabecera X-Nueva-Revision, la revisión nueva.
    """
    if not archivo.filename.lower().endswith((".docx", ".doc")):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos Word (.docx).")
    try:
        pares: List[dict] = json.loads(reemplazos or "[]")
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="El parámetro 'reemplazos' no es JSON válido.")

    datos = await archivo.read()
    try:
        doc = Document(io.BytesIO(datos))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo abrir el Word: {e}")

    # Aplicar reemplazos de contenido
    aplicados = 0
    for par in pares:
        buscar = (par or {}).get("buscar", "")
        reemplazar = (par or {}).get("reemplazar", "")
        if not buscar:
            continue
        for p in _iterar_parrafos(doc):
            if _reemplazar_en_parrafo(p, buscar, reemplazar):
                aplicados += 1

    # Subir revisión
    nueva_rev = None
    if str(subir_revision).lower() in ("true", "1", "si", "sí"):
        nueva_rev = bump_revision(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    headers = {
        "Content-Disposition": f'attachment; filename="{archivo.filename}"',
        "X-Nueva-Revision": nueva_rev or "",
        "X-Reemplazos-Aplicados": str(aplicados),
    }
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", "8000"))
    uvicorn.run(app, host="0.0.0.0", port=port)
